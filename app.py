import streamlit as st
from docx import Document
import io

st.title("英作文アプリ（大学受験向け） - Level 1")

questions = [
    ("1", "私は高校生です。", "I am a high school student."),
    ("2", "彼はとても親切な先生です。", "He is a very kind teacher."),
    ("3", "あなたは部活動に入っていますか？", "Are you in a club activity?"),
    ("4", "彼女はバスケットボール部に所属しています。", "She is in the basketball club."),
    ("5", "彼らは月曜日に英語を勉強します。", "They study English on Monday."),
    ("6", "私の兄はよく音楽を聴きます。", "My brother often listens to music."),
    ("7", "その店は8時に開きます。", "The store opens at eight."),
    ("8", "私は毎日自転車で学校に行きます。", "I go to school by bike every day."),
    ("9", "彼女は猫が好きです。", "She likes cats."),
    ("10", "あなたのお父さんは医者ですか？", "Is your father a doctor?")
]

if st.button("Wordファイルを生成する"):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("英作文アプリ：Level 1（be動詞・一般動詞）", 0)
    for q in questions:
        doc.add_paragraph(f"{q[0]}. {q[1]}")
        doc.add_paragraph("【解答】" + q[2])
        doc.add_paragraph("")
    doc.save(buffer)
    st.download_button("ダウンロード", buffer.getvalue(), file_name="level1_output.docx")
